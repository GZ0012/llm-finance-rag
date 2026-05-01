# src/loader.py

from pathlib import Path
from docx import Document


def _load_docx(file_path: str) -> list[str]:
    doc = Document(file_path)
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Tables: each row becomes a pipe-separated string so the chunker
    # treats it as a single unit rather than splitting mid-row.
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                paragraphs.append(" | ".join(row_text))

    return paragraphs


def _load_pdf(file_path: str) -> list[str]:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumber is required for PDF support. "
            "Install it with: pip install pdfplumber"
        )

    paragraphs = []

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            # Locate table bounding boxes on this page so we can exclude
            # them from prose extraction — without this, table cells bleed
            # into the text output and create duplicate/garbled paragraphs.
            table_regions = page.find_tables()

            if table_regions:
                # Crop the page to non-table area before extracting prose.
                # pdfplumber's outside_bbox() returns a new cropped page object.
                non_table_page = page
                for region in table_regions:
                    non_table_page = non_table_page.outside_bbox(region.bbox)
                prose = non_table_page.extract_text()
            else:
                prose = page.extract_text()

            # Split prose into lines; skip very short lines (page numbers,
            # stray characters) that add noise without semantic value.
            if prose:
                for line in prose.split("\n"):
                    line = line.strip()
                    if len(line) > 15:
                        paragraphs.append(line)

            # Extract tables as pipe-separated rows — same format as docx
            # tables so the rest of the pipeline needs no changes.
            for table in page.extract_tables():
                for row in table:
                    # Cells can be None when pdfplumber finds an empty cell.
                    cells = [str(c).strip() for c in row if c and str(c).strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))

    return paragraphs


def _load_txt(file_path: str) -> list[str]:
    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    return [line.strip() for line in lines if line.strip()]


def load_document(file_path: str) -> list[str]:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    # Dispatch to the right loader based on file type.
    if suffix == ".docx":
        return _load_docx(file_path)
    elif suffix == ".pdf":
        return _load_pdf(file_path)
    elif suffix in {".txt", ".md"}:
        return _load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type '{suffix}'. Supported: .docx, .pdf, .txt, .md")


